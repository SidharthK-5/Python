"""
Creating Word files using python-docx library
"""

from docx import Document
from docx.shared import Inches

# create a Document object
document = Document()

# adding heading style
document.add_heading("Docx File Write", 0)

p = document.add_paragraph("Paragraph 1 ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic.").italic = True

document.add_heading("Heading, level 1", level=1)
document.add_paragraph("Paragraph with Intense quote", style="Intense Quote")

document.add_paragraph("Unordered list", style="List Bullet")
document.add_paragraph("Ordered list", style="List Number")

document.add_heading("Sample Picture", level=2)
# adding picture into the document
document.add_picture("data/naruto.jpg", width=Inches(5))

# creating table
records = ((1, "101", "Apple"), (2, "102", "Orange"), (3, "103", "Mango"))

document.add_heading("Sample Table", level=2)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "Id"
hdr_cells[2].text = "Desc"
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save("data/generated_file.docx")
